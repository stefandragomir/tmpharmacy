
import os
import sys
import random
import base64

try:
    _path = os.path.split(os.path.split(os.path.abspath(__file__))[0])[0]
    sys.path.append(_path)
except:
    pass

_path = os.path.split(os.path.split(os.path.abspath("__file__"))[0])[0]
sys.path.append(_path)

from math                         import floor
from docx                         import Document
from docx.enum.text               import WD_ALIGN_PARAGRAPH
from PyQt5.QtCore                 import *
from PyQt5.QtGui                  import *
from PyQt5.QtWidgets              import * 
from pharm_widgets.pharm_widgets  import *
from pharm_widgets.pharm_css      import *
from pharm_icons.pharm_icons      import Pharm_Icon
from pharm_icons.pharm_icons      import Pharm_Pixmap
from datetime                     import datetime
from datetime                     import timedelta
from pharm_db.pharm_db            import PHARM_DB
from docx                         import Document
from docx.enum.text               import WD_BREAK
from datetime                     import datetime
from functools                    import partial
from docx.oxml                    import OxmlElement
from docx.oxml.ns                 import qn
from io                           import StringIO
from io                           import BytesIO
from pprint                       import pprint
from pharm_stats.pharm_stats      import Pharm_Stats_DB

"""*************************************************************************************************
****************************************************************************************************
*************************************************************************************************"""

PHARM_NUMBER_OF_TEST_QUESTIONS = 10
PHARM_MIN_CORECT_QUESTIONS     = 8
PHARM_ANSWER_OPTIONS           = ["a","b","c","d","e","f","g","h"]
PHARM_QUESTIONS_HEADING        = """
Intrebari Examen Admitere 
Universitatea de Medicină şi Farmacie Victor Babeş 
Timişoara
""" 
PHARM_TEST_HEADING        = """
Intrebari Examen Admitere 
Universitatea de Medicină şi Farmacie Victor Babeş 
Timişoara
""" 

"""*************************************************************************************************
****************************************************************************************************
*************************************************************************************************"""
class Pharm_UI(QMainWindow):

    def __init__(self,db):

        QMainWindow.__init__(self)

        self.stack_index  = 0
        self.db           = db   
        self.dekstops     = []   
        self.stats        = None

        self.create_stats()

        self.draw_gui() 

    def draw_gui(self):

        self.setWindowTitle("Admitere Farmacie 1.1.0")
        self.setWindowIcon(Pharm_Icon("pharm"))
        self.setMinimumSize(1300, 800)       
        self.setMinimumHeight(800)

        self.setStyleSheet(PHARM_CSS)

        self.wdg_central = QWidget()
        
        #selection tree
        self.tree = Pharm_WDG_Tree()
        self.tree.setFixedWidth(280)
        self.populate_tree()
        self.tree.currentItemChanged.connect(self.tree_select)

        self.toolbar_layout = QHBoxLayout()

        #left tree layout area
        self.tree_area = QVBoxLayout()   
        self.tree_area.addLayout(self.toolbar_layout) 
        self.tree_area.addWidget(self.tree)   

        #context stack
        self.context = Pharm_WDG_Stack()
        self.populate_stack()

        #Layout       
        self.top_layout = QHBoxLayout()
        self.top_layout.addLayout(self.tree_area)
        self.top_layout.addWidget(self.context)

        self.main_layout = QVBoxLayout()
        self.main_layout.addLayout(self.top_layout)

        self.wdg_central.setLayout(self.main_layout)  

        self.setCentralWidget(self.wdg_central)
        self.activateWindow() 

    def showEvent(self, event):    

        pass
        
    def tree_select(self, current, previous):

        self.context.setCurrentIndex(0) 

        _item_cfg = str(current.data(1,Qt.UserRole))

        #set context widget based on coresponding state     
        self.context.setCurrentIndex(int(_item_cfg))  

    def populate_tree(self):

        self.tree.clear()

        self.stack_index = 1 

        _parent = self.tree.invisibleRootItem()

        for _category in self.db:

            _item = QTreeWidgetItem(_parent)
            _item.setData(0, Qt.EditRole, _category.name)             

            _serial_cfg = self.stack_index                    
            _item.setData(1, Qt.UserRole, _serial_cfg)

            _item.setIcon(0,Pharm_Icon("test"))

            self.stack_index += 1

    def populate_stack(self):

        self.context.insertWidget(0, QWidget())

        for _idx in range(len(self.db)):

            _desktop = Pharm_WDG_Desktop(self.db[_idx],self.stats)

            self.dekstops.append(_desktop)

            self.context.insertWidget(_idx + 1, _desktop) 

        self.context.setCurrentIndex(0)  

    def create_stats(self):

        _path = os.path.join(os.path.expanduser("~"),".pharm")

        if not os.path.exists(_path):

            os.mkdir(_path)

        _path = os.path.join(_path,"pharm.db")

        self.stats = Pharm_Stats_DB(_path)

"""*************************************************************************************************
****************************************************************************************************
*************************************************************************************************"""
class Pharm_Model_Test(object):

    def __init__(self,stats,category):

        self.questions = []
        self.stats     = stats
        self.category  = category

    def clear(self):

        for _question in self.questions:

            for _answer in _question.answers:

                _answer.is_selected = False

    def get_score(self,question):

        _user_scor = 0
        _nrm_scor  = 0

        for _answer in question.answers:

            if _answer.corect and  _answer.selected:
                _user_scor += 1

            if _answer.corect:
                _nrm_scor += 1

        return _user_scor,_nrm_scor

    def get_result(self):

        _total    = len(self.questions)
        _corect   = 0
        _incorect = 0

        for _question in self.questions:

            _user_scor,_nrm_scor = self.get_score(_question)

            if _user_scor == _nrm_scor:
                _corect += 1
            else:
                _incorect += 1

            self.add_stats_question(_question,1)

        return _corect,_incorect

    def add_stats_question(self,question,type):

        _status    = 0

        _user_scor,_nrm_scor = self.get_score(question)

        if _user_scor == _nrm_scor:
            _status = 1
        else:
            _status = 0

        self.stats.add_question(self.category.name,_status)

"""*************************************************************************************************
****************************************************************************************************
*************************************************************************************************"""
class Pharm_WDG_Desktop(QWidget):

    def __init__(self,category,stats):

        QWidget.__init__(self)

        self.category = category
        self.stats    = stats

        self.draw_gui()

    def draw_gui(self):

        self.bt_test  = Pharm_WDG_Button("Test",              Pharm_Icon("test_normal"),   Pharm_Icon("test_hover"),    "#606060")
        self.bt_learn = Pharm_WDG_Button("Invata",            Pharm_Icon("learn_normal"),  Pharm_Icon("learn_hover"),   "#606060")
        self.bt_gen   = Pharm_WDG_Button("Genereaza",         Pharm_Icon("generate_test"), Pharm_Icon("generate_test"), "#606060")
        self.bt_exp   = Pharm_WDG_Button("Exporta",           Pharm_Icon("generate_doc"),  Pharm_Icon("generate_doc"),  "#606060")
        self.bt_stats = Pharm_WDG_Button("Statistici",        Pharm_Icon("stats"),         Pharm_Icon("stats"),         "#606060")

        self.bt_test.setIconSize(QSize(100,100))
        self.bt_learn.setIconSize(QSize(100,100))
        self.bt_gen.setIconSize(QSize(100,100))
        self.bt_exp.setIconSize(QSize(100,100))
        self.bt_stats.setIconSize(QSize(100,100))

        self.bt_test.clicked.connect(self.clbk_bt_test)
        self.bt_learn.clicked.connect(self.clbk_bt_learn)
        self.bt_gen.clicked.connect(self.clbk_bt_gen)
        self.bt_exp.clicked.connect(self.clbk_bt_exp)
        self.bt_stats.clicked.connect(self.clbk_bt_stats)

        self.wdg_test  = Pharm_WDG_Desktop_Test(self,self.category,self.stats)
        self.wdg_stats = Pharm_WDG_Desktop_Stats(self)

        self.bt_layout = QHBoxLayout()
        self.bt_layout.addWidget(self.bt_learn)
        self.bt_layout.addWidget(self.bt_test)
        self.bt_layout.addWidget(self.bt_gen)
        self.bt_layout.addWidget(self.bt_exp)
        self.bt_layout.addWidget(self.bt_stats)

        self.main_layout = QVBoxLayout()
        self.main_layout.addLayout(self.bt_layout)
        self.main_layout.addWidget(self.wdg_test)
        self.main_layout.addWidget(self.wdg_stats)

        self.wdg_test.hide()

        self.wdg_stats.hide()

        self.setLayout(self.main_layout)  

    def get_docs_path(self):

        _path = os.path.abspath(__file__)
        _path = os.path.split(_path)[0]
        _path = os.path.join(_path,"docs")

        if not os.path.exists(_path):

            os.mkdir(_path)

        return _path

    def get_generated_doc_path(self,name):

        _path = self.get_docs_path()

        _timestamp = str(datetime.now())

        _file_name = "intrebari_%s_%s.docx" % (name,_timestamp.replace(":","_").replace(" ","_").replace("/","_").replace("-","_").replace(".","_"))

        _path = os.path.join(_path,_file_name)

        return _path,_timestamp

    def get_generated_test_path(self,name,timestamp=None):

        _path = self.get_docs_path()

        if timestamp == None:
            _timestamp = str(datetime.now())
        else:
            _timestamp = timestamp
            _path = os.path.join(_path,_timestamp.replace(":","_").replace(" ","_").replace("/","_").replace("-","_").replace(".","_"))

            if not os.path.exists(_path):

                os.mkdir(_path)

        _file_name = "test_%s_%s.docx" % (name,_timestamp.replace(":","_").replace(" ","_").replace("/","_").replace("-","_").replace(".","_"))

        _path = os.path.join(_path,_file_name)

        return _path,_timestamp

    def hide_buttons(self):

        self.bt_test.hide()

        self.bt_learn.hide()

        self.bt_gen.hide()

        self.bt_stats.hide()

        self.bt_exp.hide()

    def clbk_bt_test(self,state):

        self.wdg_test.show()

        self.wdg_stats.hide()

        self.hide_buttons()

        self.wdg_test.start("test")

    def clbk_bt_learn(self,state):

        self.wdg_test.show()

        self.wdg_stats.hide()

        self.hide_buttons()

        self.wdg_test.start("learn")

    def clbk_bt_gen(self,state,timestamp=None):

        _questions_indexes = random.sample(range(len(self.category.questions)), PHARM_NUMBER_OF_TEST_QUESTIONS)

        _path, _timestamp = self.get_generated_test_path(self.category.name,timestamp)

        _document = Document()

        _paragraph_format             = _document.styles['Normal'].paragraph_format
        _paragraph_format.space_after = 3

        self.__generate_top_table(_document)

        _paragraph = _document.add_paragraph()
        _paragraph.add_run("%s\n%s" % (PHARM_TEST_HEADING,self.category.name.upper())).bold = True
        _paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _document.add_paragraph("Test generat la data: %s" % (_timestamp,))

        _question_count = 0

        _table_questions = []

        for _index in range(len(self.category.questions)):

            if _index in _questions_indexes:

                _question = self.category.questions[_index]

                _table_questions.append([_index + 1,_question])

                _question_text = "%s. %s" % (_question_count + 1,_question.text)

                _question_count += 1

                _paragraph = _document.add_paragraph(_question_text)

                if _question.image != None:

                    _bytes = base64.b64decode(_question.image)

                    _str = BytesIO(_bytes)

                    _document.add_picture(_str)

                _answer_count = 0

                for _answer in _question.answers:

                    _answer_text = "    %s) %s" % (PHARM_ANSWER_OPTIONS[_answer_count],_answer.text)

                    _answer_count += 1

                    _paragraph = _document.add_paragraph(_answer_text)

        _document.add_paragraph("Test generat la data: %s" % (_timestamp,))

        self.__generate_bottom_table(_document,_table_questions)

        _document.add_paragraph("Test generat la data: %s" % (_timestamp,))

        _document.save(_path)

        if timestamp == None:

            _msg = QMessageBox()
            _msg.setWindowIcon(Pharm_Icon("pharm"))
            _msg.setIcon(QMessageBox.Information)
            _msg.setText("Generat chestionar test in fisierul\n%s" % (os.path.split(_path)[1],))
            _msg.setWindowTitle("Generat Chestionar Test %s" % (self.category.name.upper(),))
            _msg.exec_()

        
            os.startfile(_path)

        return _path

    def clbk_bt_stats(self,state):

        self.wdg_test.hide()

        self.wdg_stats.show()

        self.hide_buttons()

        self.wdg_stats.refresh()

    def __generate_top_table(self,document):

            _table = document.add_table(rows=2, cols=3)

            _table.cell(0, 0).text = "NUME SI PRENUME "
            self.set_cell_border(
                                    _table.cell(0, 0),
                                    top={"sz": 5, "val": "single", "color": "#000000", "space": "0"},
                                    bottom={"sz": 5, "val": "single", "color": "#000000", "space": "0"},
                                    start={"sz": 5, "val": "single", "color": "#000000", "space": "0"},
                                    end={"sz": 5, "val": "single", "color": "#000000", "space": "0"})

            _table.cell(0, 1).text = "SEMNATURA"
            self.set_cell_border(
                                    _table.cell(0, 1),
                                    top={"sz": 5, "val": "single", "color": "#000000", "space": "0"},
                                    bottom={"sz": 5, "val": "single", "color": "#000000", "space": "0"},
                                    start={"sz": 5, "val": "single", "color": "#000000", "space": "0"},
                                    end={"sz": 5, "val": "single", "color": "#000000", "space": "0"})

            _table.cell(0, 2).text = "DATA"
            self.set_cell_border(
                                    _table.cell(0, 2),
                                    top={"sz": 5, "val": "single", "color": "#000000", "space": "0"},
                                    bottom={"sz": 5, "val": "single", "color": "#000000", "space": "0"},
                                    start={"sz": 5, "val": "single", "color": "#000000", "space": "0"},
                                    end={"sz": 5, "val": "single", "color": "#000000", "space": "0"})

            self.set_cell_border(
                                    _table.cell(1, 0),
                                    top={"sz": 5, "val": "single", "color": "#000000", "space": "0"},
                                    bottom={"sz": 5, "val": "single", "color": "#000000", "space": "0"},
                                    start={"sz": 5, "val": "single", "color": "#000000", "space": "0"},
                                    end={"sz": 5, "val": "single", "color": "#000000", "space": "0"})

            self.set_cell_border(
                                    _table.cell(1, 1),
                                    top={"sz": 5, "val": "single", "color": "#000000", "space": "0"},
                                    bottom={"sz": 5, "val": "single", "color": "#000000", "space": "0"},
                                    start={"sz": 5, "val": "single", "color": "#000000", "space": "0"},
                                    end={"sz": 5, "val": "single", "color": "#000000", "space": "0"})

            self.set_cell_border(
                                    _table.cell(1, 2),
                                    top={"sz": 5, "val": "single", "color": "#000000", "space": "0"},
                                    bottom={"sz": 5, "val": "single", "color": "#000000", "space": "0"},
                                    start={"sz": 5, "val": "single", "color": "#000000", "space": "0"},
                                    end={"sz": 5, "val": "single", "color": "#000000", "space": "0"})

    def __generate_bottom_table(self,document,questions):

            document.add_page_break()

            _paragraph = document.add_paragraph()
            _paragraph.add_run("%s\n%s" % (PHARM_TEST_HEADING,self.category.name.upper())).bold = True
            _paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

            _table = document.add_table(rows=11, cols=4)

            _table.cell(0, 0).text = "Numar Intrebare Test"
            self.set_cell_border(
                                    _table.cell(0, 0),
                                    top={"sz": 5, "val": "single", "color": "#000000", "space": "0"},
                                    bottom={"sz": 5, "val": "single", "color": "#000000", "space": "0"},
                                    start={"sz": 5, "val": "single", "color": "#000000", "space": "0"},
                                    end={"sz": 5, "val": "single", "color": "#000000", "space": "0"})

            _table.cell(0, 1).text = "Numar Intrebare Documentatie"
            self.set_cell_border(
                                    _table.cell(0, 1),
                                    top={"sz": 5, "val": "single", "color": "#000000", "space": "0"},
                                    bottom={"sz": 5, "val": "single", "color": "#000000", "space": "0"},
                                    start={"sz": 5, "val": "single", "color": "#000000", "space": "0"},
                                    end={"sz": 5, "val": "single", "color": "#000000", "space": "0"})

            _table.cell(0, 2).text = "Raspuns Corect"
            self.set_cell_border(
                                    _table.cell(0, 2),
                                    top={"sz": 5, "val": "single", "color": "#000000", "space": "0"},
                                    bottom={"sz": 5, "val": "single", "color": "#000000", "space": "0"},
                                    start={"sz": 5, "val": "single", "color": "#000000", "space": "0"},
                                    end={"sz": 5, "val": "single", "color": "#000000", "space": "0"})

            _table.cell(0, 3).text = "Raspuns Examen"
            self.set_cell_border(
                                    _table.cell(0, 3),
                                    top={"sz": 5, "val": "single", "color": "#000000", "space": "0"},
                                    bottom={"sz": 5, "val": "single", "color": "#000000", "space": "0"},
                                    start={"sz": 5, "val": "single", "color": "#000000", "space": "0"},
                                    end={"sz": 5, "val": "single", "color": "#000000", "space": "0"})

            for _index in range(len(questions)):

                _table.cell(_index + 1, 0).text = str(_index + 1)  

                self.set_cell_border(
                                        _table.cell(_index + 1, 0),
                                        top={"sz": 5, "val": "single", "color": "#000000", "space": "0"},
                                        bottom={"sz": 5, "val": "single", "color": "#000000", "space": "0"},
                                        start={"sz": 5, "val": "single", "color": "#000000", "space": "0"},
                                        end={"sz": 5, "val": "single", "color": "#000000", "space": "0"})

            for _index in range(len(questions)):

                _table.cell(_index + 1, 1).text = str(questions[_index][0])  

                self.set_cell_border(
                                        _table.cell(_index + 1, 1),
                                        top={"sz": 5, "val": "single", "color": "#000000", "space": "0"},
                                        bottom={"sz": 5, "val": "single", "color": "#000000", "space": "0"},
                                        start={"sz": 5, "val": "single", "color": "#000000", "space": "0"},
                                        end={"sz": 5, "val": "single", "color": "#000000", "space": "0"})

            for _index in range(len(questions)):

                _answer_text = ""

                _answer_count = 0

                for _answer in questions[_index][1].answers:

                    if _answer.corect:

                        _answer_text += "%s, " % (PHARM_ANSWER_OPTIONS[_answer_count],)

                    _answer_count += 1


                _table.cell(_index + 1, 2).text = _answer_text

                self.set_cell_border(
                                        _table.cell(_index + 1, 2),
                                        top={"sz": 5, "val": "single", "color": "#000000", "space": "0"},
                                        bottom={"sz": 5, "val": "single", "color": "#000000", "space": "0"},
                                        start={"sz": 5, "val": "single", "color": "#000000", "space": "0"},
                                        end={"sz": 5, "val": "single", "color": "#000000", "space": "0"})

            for _index in range(len(questions)):

                self.set_cell_border(
                                        _table.cell(_index + 1, 3),
                                        top={"sz": 5, "val": "single", "color": "#000000", "space": "0"},
                                        bottom={"sz": 5, "val": "single", "color": "#000000", "space": "0"},
                                        start={"sz": 5, "val": "single", "color": "#000000", "space": "0"},
                                        end={"sz": 5, "val": "single", "color": "#000000", "space": "0"})

    def set_cell_border(sefl, cell, **kwargs):

        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()

        # check for tag existnace, if none found, then create one
        tcBorders = tcPr.first_child_found_in("w:tcBorders")
        if tcBorders is None:
            tcBorders = OxmlElement('w:tcBorders')
            tcPr.append(tcBorders)

        # list over all available tags
        for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
            edge_data = kwargs.get(edge)
            if edge_data:
                tag = 'w:{}'.format(edge)

                # check for tag existnace, if none found, then create one
                element = tcBorders.find(qn(tag))
                if element is None:
                    element = OxmlElement(tag)
                    tcBorders.append(element)

                # looks like order of attributes is important
                for key in ["sz", "val", "color", "space", "shadow"]:
                    if key in edge_data:
                        element.set(qn('w:{}'.format(key)), str(edge_data[key]))

    def clbk_bt_exp(self,state):

        _path, _timestamp = self.get_generated_doc_path(self.category.name)

        _document = Document()

        _paragraph = _document.add_paragraph()
        _paragraph.add_run("%s\n%s" % (PHARM_QUESTIONS_HEADING,self.category.name.upper())).bold = True
        _paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _document.add_paragraph("")
        _document.add_paragraph("")

        _question_count = 1

        for _question in self.category.questions:

            _question_text = "%s. %s" % (_question_count,_question.text)

            _question_count += 1

            _paragraph = _document.add_paragraph(_question_text)

            if _question.image != None:

                _bytes = base64.b64decode(_question.image)

                _str = BytesIO(_bytes)

                _document.add_picture(_str)

            _answer_count = 0


            for _answer in _question.answers:

                _answer_text = "    %s) %s" % (PHARM_ANSWER_OPTIONS[_answer_count],_answer.text)

                _answer_count += 1

                _paragraph = _document.add_paragraph()
                _run       = _paragraph.add_run(_answer_text)
                _run.bold  = _answer.corect

            _document.add_paragraph("")

        _document.add_paragraph("")
        _document.add_paragraph("")
        _document.add_paragraph("")
        _document.add_paragraph("")
        _document.add_paragraph("Generat la data: %s" % (_timestamp,))

        _document.save(_path)

        _msg = QMessageBox()
        _msg.setWindowIcon(Pharm_Icon("pharm"))
        _msg.setIcon(QMessageBox.Information)
        _msg.setText("Generat intrebari in fisierul\n%s" % (os.path.split(_path)[1],))
        _msg.setWindowTitle("Generat Intrebari %s" % (self.category.name.upper(),))
        _msg.exec_()

        os.startfile(_path)

"""*************************************************************************************************
****************************************************************************************************
*************************************************************************************************"""
class Pharm_WDG_Desktop_Stats(QWidget):

    def __init__(self,parent):

        QWidget.__init__(self)

        self.parent = parent
        self.stats  = parent.stats

        self.draw_gui()

    def draw_gui(self):

        self.bt_close           = Pharm_WDG_Small_Button( Pharm_Icon("close_normal"),      Pharm_Icon("close_hover"),    self.clbk_close, "Inchide")
        self.wdg_plot_questions = Pharm_WDG_Plot_StackedBar_Periods(self)
        self.wdg_plot_tests     = Pharm_WDG_Plot_StackedBar_Periods(self)
        self.wdg_plot_time      = Pharm_WDG_Plot_StackedBar_Periods(self)

        self.bt_close.setIconSize(QSize(30,30))

        self.main_layout = QVBoxLayout()
        
        self.main_layout.addWidget(self.bt_close)
        self.main_layout.addWidget(self.wdg_plot_questions.chart_view)
        self.main_layout.addWidget(self.wdg_plot_tests.chart_view)
        self.main_layout.addWidget(self.wdg_plot_time.chart_view)

        self.main_layout.setAlignment(Qt.AlignTop)

        self.setLayout(self.main_layout) 

        self.refresh()

    def clbk_close(self,state):

        self.parent.bt_test.show()
        self.parent.bt_learn.show()
        self.parent.bt_gen.show()
        self.parent.bt_stats.show()
        self.parent.bt_exp.show()
        self.parent.wdg_test.hide()
        self.parent.wdg_stats.hide()
    
    def get_last_days(self,days):

        _periods = []

        _last_day  = datetime.now()
        _first_day = _last_day - timedelta(days=days)

        for _index in range(days + 1):

            _day = _first_day + timedelta(days=_index)

            _periods.append([
                                "%s-%s-%s" % (_day.year,_day.month,_day.day),
                                _day.replace(hour=0,minute=0,second=0),
                                _day.replace(hour=23,minute=59,second=59),
                            ])

        return _periods      
    
    def read_questions(self):

        _values = [[],[]]
        _labels = []

        _periods = self.get_last_days(15)

        _questions = self.stats.read_questions()

        for _period in _periods:

            _status_ok  = 0
            _status_nok = 0

            _labels.append(_period[0])

            for _question in _questions:

                _time = datetime.strptime(_question.time,"%Y-%m-%d %H:%M:%S.%f")

                if _time >= _period[1] and _time <= _period[2]:
                    if 1 == _question.status:
                        _status_ok += 1
                    else:
                        _status_nok += 1

            _values[0].append(_status_ok)
            _values[1].append(_status_nok)

        return _values, _labels

    def read_tests(self):

        _values = [[],[]]
        _labels = []

        _periods = self.get_last_days(15)

        _tests = self.stats.read_tests()

        for _period in _periods:

            _status_ok  = 0
            _status_nok = 0

            _labels.append(_period[0])

            for _test in _tests:

                _time = datetime.strptime(_test.time,"%Y-%m-%d %H:%M:%S.%f")

                if _time >= _period[1] and _time <= _period[2]:
                    if 1 == _test.status:
                        _status_ok += 1
                    else:
                        _status_nok += 1

            _values[0].append(_status_ok)
            _values[1].append(_status_nok)

        return _values, _labels

    def read_time(self):

        _values = [[]]
        _labels = []

        _periods = self.get_last_days(15)

        _tests = self.stats.read_tests()

        for _period in _periods:

            _duration  = 0

            _labels.append(_period[0])

            for _test in _tests:

                _time = datetime.strptime(_test.time,"%Y-%m-%d %H:%M:%S.%f")

                if _time >= _period[1] and _time <= _period[2]:
                    _duration += _test.duration

            _values[0].append(floor(_duration/60))

        return _values, _labels


    def refresh(self):

        _values, _x_labels = self.read_questions()

        self.wdg_plot_questions.clear()

        _y_labels        = ["Raspunsuri Corecte","Raspunsuri Incorecte"]
        _y_labels_colors = ["#1da14f","#ba1c34"]

        self.wdg_plot_questions.draw(
                                        "",
                                        _values,
                                        _y_labels,
                                        _y_labels_colors,
                                        _x_labels)

        _values, _x_labels = self.read_tests()

        self.wdg_plot_tests.clear()

        _y_labels        = ["Teste Trecute","Teste Picate"]
        _y_labels_colors = ["#1da14f","#ba1c34"]

        self.wdg_plot_tests.draw(
                                    "",
                                    _values,
                                    _y_labels,
                                    _y_labels_colors,
                                    _x_labels)

        _values, _x_labels = self.read_time()

        self.wdg_plot_time.clear()

        _y_labels        = ["Minute in Test"]
        _y_labels_colors = ["#4b27e8"]

        self.wdg_plot_time.draw(
                                    "",
                                    _values,
                                    _y_labels,
                                    _y_labels_colors,
                                    _x_labels)

"""*************************************************************************************************
****************************************************************************************************
*************************************************************************************************"""
class Pharm_WDG_Desktop_Test(QWidget):

    def __init__(self,parent,category,stats):

        QWidget.__init__(self)

        self.category        = category
        self.question_number = 0
        self.parent          = parent
        self.test_type       = ""
        self.stats           = stats
        self.test_learn      = Pharm_Model_Test(self.stats,self.category)
        self.test_exam       = Pharm_Model_Test(self.stats,self.category)
        self.is_finished     = False   
        self.start_time      = datetime.now()     

        self.test_learn.questions = self.category.questions

        self.draw_gui()

    def draw_gui(self):

        self.bt_next   = Pharm_WDG_Small_Button( Pharm_Icon("next_normal"),       Pharm_Icon("next_hover"),     self.clbk_next , "Intrebare urmatoare")
        self.bt_prev   = Pharm_WDG_Small_Button( Pharm_Icon("previous_normal"),   Pharm_Icon("previous_hover"), self.clbk_prev, "Intrebarea precedenta")
        self.bt_result = Pharm_WDG_Small_Button( Pharm_Icon("results_normal"),    Pharm_Icon("result_hover"),   self.clbk_result, "Rezultat")
        self.bt_close  = Pharm_WDG_Small_Button( Pharm_Icon("close_normal"),      Pharm_Icon("close_hover"),    self.clbk_close, "Inchide")

        self.bt_next.setIconSize(QSize(50,50))
        self.bt_prev.setIconSize(QSize(50,50))
        self.bt_result.setIconSize(QSize(50,50))
        self.bt_close.setIconSize(QSize(50,50))

        self.lbl_status = Pharm_WDG_Label()
        self.lbl_result = Pharm_WDG_Label()

        self.lbl_status.setStyleSheet("font-size: 25px;")
        self.lbl_result.setStyleSheet("font-size: 25px;")

        self.bt_layout = QHBoxLayout()
        
        self.bt_layout.addWidget(self.bt_close)
        self.bt_layout.addWidget(self.bt_prev)
        self.bt_layout.addWidget(self.bt_result)
        self.bt_layout.addWidget(self.bt_next)

        self.wdg_question = Pharm_WDG_Question()

        _policy = QSizePolicy(QSizePolicy.Expanding, QSizePolicy.Expanding)
        
        self.wdg_question.setSizePolicy(_policy)

        self.main_layout = QVBoxLayout()
        
        self.main_layout.addLayout(self.bt_layout)
        self.main_layout.addWidget(self.wdg_question)
        self.main_layout.addWidget(self.lbl_status)
        self.main_layout.addWidget(self.lbl_result)

        self.main_layout.setAlignment(Qt.AlignTop)

        self.setLayout(self.main_layout) 

        self.lbl_result.hide()

        self.bt_next.hide()
        self.bt_prev.hide()
        self.bt_close.hide()
        self.bt_result.hide()

    def start(self,test_type):

        self.is_finished = False

        self.test_type = test_type

        self.question_number = 0

        self.wdg_question.show()

        if self.test_type == "learn":            
            self.test_learn.clear()
            self.wdg_question.populate(self.test_learn.questions[self.question_number], True)
            self.bt_next.show()
            self.bt_close.show()
            self.bt_result.show()
            self.bt_prev.show()
        else:            
            self.test_exam.clear()
            self.bt_close.show()
            self.bt_next.show()
            self.bt_result.hide()
            self.bt_prev.show()
            self.get_test_questions()
            self.wdg_question.populate(self.test_exam.questions[self.question_number], True)

        self.lbl_result.hide()
        self.set_status()

        self.start_time = datetime.now()     

    def get_test_questions(self):

        global PHARM_NUMBER_OF_TEST_QUESTIONS

        if PHARM_NUMBER_OF_TEST_QUESTIONS < len(self.test_learn.questions):

            _questions_indexes = random.sample(range(len(self.test_learn.questions)), PHARM_NUMBER_OF_TEST_QUESTIONS)

            self.test_exam.questions = [self.test_learn.questions[_index] for _index in _questions_indexes]
        else:
            self.test_exam.questions = self.test_learn.questions

    def clbk_next(self,state):

        for _index in range(len(self.test_learn.questions[self.question_number].answers)):
            
            self.wdg_question.rd_answers[_index].set_text_normal()

        if not self.is_end(self.question_number):

            self.question_number += 1

            if self.test_type == "learn":
                self.wdg_question.populate(self.test_learn.questions[self.question_number],True)
            else:
                self.wdg_question.populate(self.test_exam.questions[self.question_number],not self.is_finished)

            self.set_status()
        else:
            pass

        if self.test_type != "learn" and not self.is_finished:
            if self.question_number == len(self.test_exam.questions) - 1:
                self.bt_result.show()
            else:
                self.bt_result.hide()

        if self.is_finished:

            self.color_questions_status()

    def clbk_prev(self,state):

        for _index in range(len(self.test_learn.questions[self.question_number].answers)):

            self.wdg_question.rd_answers[_index].set_text_normal()

        if self.test_type != "learn":

            self.bt_result.hide()

        if not self.is_begining(self.question_number):

            self.question_number -= 1

            if self.test_type == "learn":
                self.wdg_question.populate(self.test_learn.questions[self.question_number],True)
            else:
                self.wdg_question.populate(self.test_exam.questions[self.question_number],not self.is_finished)

            self.set_status()

        else:
            pass

        if self.is_finished:

            self.color_questions_status()

    def clbk_close(self,state):

        self.parent.bt_test.show()
        self.parent.bt_learn.show()
        self.parent.bt_gen.show()
        self.parent.bt_stats.show()
        self.parent.bt_exp.show()
        self.parent.wdg_test.hide()

    def clbk_result(self,state):

        _corect   = 0
        _incorect = 0
        _total    = 0

        if self.test_type == "learn":

            self.color_questions_status()

            self.test_learn.add_stats_question(self.test_learn.questions[self.question_number],0)
        else:
            self.is_finished = True

            self.wdg_question.populate(self.test_exam.questions[self.question_number],not self.is_finished)

            self.color_questions_status()

            _corect,_incorect = self.test_exam.get_result()
            self.bt_result.hide()
            self.bt_close.show()

            self.lbl_result.show()

            if _corect >= PHARM_MIN_CORECT_QUESTIONS:
                self.lbl_result.setText("ADMIS Corecte[%s] Incorecte[%s]" % (_corect,_incorect))
                self.lbl_result.setStyleSheet("QLabel { background-color : #14c941; font: 18pt; color: #ffffff}")
                self.lbl_result.setAlignment(Qt.AlignCenter)
            else:
                self.lbl_result.setText("PICAT Corecte[%s] Incorecte[%s]" % (_corect,_incorect))
                self.lbl_result.setStyleSheet("QLabel { background-color : #ba2012; font: 18pt;  color: #ffffff}")
                self.lbl_result.setAlignment(Qt.AlignCenter)

            _end_time = datetime.now()

            self.stats.add_test(
                                (_end_time - self.start_time).total_seconds(),
                                _corect,
                                _incorect,
                                int(_corect >= PHARM_MIN_CORECT_QUESTIONS)
                                )

    def color_questions_status(self):

        if self.test_type == "learn":
            _question    = self.test_learn.questions[self.question_number]
        else:
            _question = self.test_exam.questions[self.question_number]

        for _index in range(len(_question.answers)):

            self.wdg_question.rd_answers[_index].set_check_state(_question.answers[_index].selected)

            if _question.answers[_index].corect:

                self.wdg_question.rd_answers[_index].set_text_corect()               

            else:
                self.wdg_question.rd_answers[_index].set_text_incorrect()

    def set_status(self):

        if self.test_type == "learn":
            self.lbl_status.setText("Intrebarea[%s/%s] " % (self.question_number + 1,len(self.test_learn.questions)))
        else:
            self.lbl_status.setText("Intrebarea[%s/%s]" % (self.question_number + 1,len(self.test_exam.questions)))

    def is_end(self,question_number):

        _state = False

        if self.test_type == "learn":
            _state = (question_number + 1) >= len(self.test_learn.questions)
        else:
            _state = (question_number + 1) >= len(self.test_exam.questions)

        return _state

    def is_begining(self,question_number):

        return question_number <= 0

"""*************************************************************************************************
****************************************************************************************************
*************************************************************************************************"""
class Pharm_WDG_Question(QWidget):

    def __init__(self):

        QWidget.__init__(self)

        self.draw_gui()

        self.question = None

    def draw_gui(self):

        self.lbl_question = Pharm_WDG_Label()
        self.lbl_question.setWordWrap(True)
        self.rd_answers   = []

        self.lbl_question.setStyleSheet("font-size: 23px;")

        for _index in range(5):
            self.rd_answers.append(Pharm_WDG_CheckBox(""))
            self.rd_answers[-1].register_checkbox_clbk(partial(self.clbk_answer,_index))

        self.main_layout = QVBoxLayout()

        self.main_layout.addWidget(self.lbl_question)

        for _index in range(5):

            self.main_layout.addWidget(self.rd_answers[_index])

        self.setLayout(self.main_layout)

        self.lbl_question.hide()

        for _index in range(5):

            self.rd_answers[_index].hide()

    def populate(self,question,with_checks):

        for _index in range(5):

            self.rd_answers[_index].hide()

        self.question = question

        self.lbl_question.show()

        for _index in range(len(self.question.answers)):

            self.rd_answers[_index].set_text_normal()

            if not with_checks:
                self.rd_answers[_index].hide_check()
            else:
                self.rd_answers[_index].show_check()

            self.rd_answers[_index].show()
            self.rd_answers[_index].set_text(self.question.answers[_index].text)

            if with_checks:
                if self.question.answers[_index].selected:
                    self.rd_answers[_index].set_check_state(True)
                else:
                    self.rd_answers[_index].set_check_state(False)

        self.lbl_question.setText(self.question.text)

    def clbk_answer(self,index,state):
        
        self.question.answers[index].selected = state == Qt.Checked

"""*************************************************************************************************
****************************************************************************************************
*************************************************************************************************"""
class Pharm(object):

    def __init__(self):

        global PHARM_DB

        self.app = QApplication(sys.argv)  

        _ui  = Pharm_UI(PHARM_DB)   

        _ui.show()

        sys.exit(self.app.exec_())

"""*************************************************************************************************
****************************************************************************************************
*************************************************************************************************"""
if __name__ == "__main__":

    Pharm()

